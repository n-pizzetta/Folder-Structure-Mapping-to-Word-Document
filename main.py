import os
from docx import Document
from docx.shared import Pt, RGBColor
from tqdm import tqdm

def collect_folder_structure(folder_path):
    """
    Collect the structure of the folder including subdirectories and files.
    """
    folder_structure = {}
    total_items = sum([len(files) + len(dirs) for _, dirs, files in os.walk(folder_path)])
    
    with tqdm(total=total_items, desc="Collecting folder structure", unit="item") as pbar:
        for root, dirs, files in os.walk(folder_path):
            folder_structure[root] = {'dirs': dirs, 'files': files}
            pbar.update(len(files) + len(dirs))
    
    return folder_structure

def get_color_by_extension(extension):
    """
    Get color based on file extension.
    """
    color_map = {
        '.txt': RGBColor(255, 0, 0),      # Red
        '.py': RGBColor(0, 128, 0),       # Green
        '.jpg': RGBColor(255, 255, 0),    # Yellow
        '.png': RGBColor(255, 165, 0),    # Orange
        '.pdf': RGBColor(128, 0, 128),    # Purple
        # Add more file types and colors as needed
    }
    return color_map.get(extension, RGBColor(0, 0, 0))  # Default to black

def add_items_to_doc(doc, folder_structure, folder, level=1):
    """
    Add folder and file items to the document with proper heading and indentation.
    """
    folder_name = os.path.basename(folder) or folder
    doc.add_heading(folder_name, level=min(level, 9))  # Ensure heading level is within the valid range
    
    # Add files
    for file_name in folder_structure[folder]['files']:
        file_extension = os.path.splitext(file_name)[1]
        color = get_color_by_extension(file_extension)
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(file_name)
        run.font.size = Pt(14 - min(level, 9))  # Adjust file name size based on folder level
        run.font.color.rgb = color
        paragraph.paragraph_format.left_indent = Pt(level * 20)
    
    # Add subdirectories
    for dir_name in folder_structure[folder]['dirs']:
        dir_path = os.path.join(folder, dir_name)
        add_items_to_doc(doc, folder_structure, dir_path, level + 1)

def create_doc(folder_structure, doc_path):
    """
    Create a Word document with the folder structure.
    """
    doc = Document()
    doc.add_heading('Network Folder Structure', level=1)
    
    root_folder = list(folder_structure.keys())[0]
    add_items_to_doc(doc, folder_structure, root_folder)

    doc.save(doc_path)

def main(folder_path, output_doc_path):
    """
    Main function to collect folder structure and create the document.
    """
    folder_structure = collect_folder_structure(folder_path)
    create_doc(folder_structure, output_doc_path)

if __name__ == "__main__":
    # Default input and output paths
    default_folder_path = r"./input_folder"
    default_output_doc_path = r"./output_folder_structure.docx"
    main(default_folder_path, default_output_doc_path)

    # For command-line usage, uncomment the following lines:
    # import argparse
    # parser = argparse.ArgumentParser(description="Generate a Word document of a folder structure.")
    # parser.add_argument("folder_path", help="Path to the folder to diagram.")
    # parser.add_argument("output_doc_path", help="Path to save the output Word document file.")
    # args = parser.parse_args()
    # main(args.folder_path, args.output_doc_path)
